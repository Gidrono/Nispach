from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
#need to change path
path = r'C:\Users\User\Desktop\python\mypython\newfold\''
number_of_files = int(input('כמה נספחים יש? '))
filesnum = list(range(1, (number_of_files+1)))
index = 1
for file in filesnum:
    filename = str(input('מה יש בנספח '+str(file)+'? '))
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.alignment = 1
    style = document.styles['Normal']
    font = style.font
    p=document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.alignment = 1
    p.add_run('נספח ' + str(file)).bold = True
    c=document.add_paragraph(filename)
    c.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.alignment = 1
    font.name = 'David'
    font.size = Pt(64)
    filex = 'nispach '+str(file)+'.docx'
    filepath = path + filex
    document.save(filepath)

import os
import win32com.client
import re
path = (r'C:\Users\User\Desktop\python\mypython\newfold')
word_file_names = []
word = win32com.client.Dispatch('Word.Application')
for dirpath, dirnames, filenames in os.walk(path):
    for f in filenames:  
        if f.lower().endswith(".docx") :
            new_name = f.replace(".docx", ".pdf")
            in_file =(dirpath + '/'+ f)
            new_file =(dirpath + '/' + new_name)
            doc = word.Documents.Open(in_file)
            doc.SaveAs(new_file, FileFormat = 17)
            doc.Close()
        if f.lower().endswith(".doc"):
            new_name = f.replace(".doc", ".pdf")
            in_file =(dirpath +'/' + f)
            new_file =(dirpath +'/' + new_name)
            doc = word.Documents.Open(in_file)
            doc.SaveAs(new_file, FileFormat = 17)
            doc.Close()
word.Quit()